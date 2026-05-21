# =============================
# AI Consult Verslag Generator
# =============================
# Vereisten:
# pip install openai python-docx

import os

print(os.listdir())

from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
from faster_whisper import WhisperModel

whisper_model = WhisperModel("base")

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

if not os.getenv("OPENAI_API_KEY"):
    raise ValueError("API key ontbreekt")

# =============================
# HELPERS
# =============================
def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def save_to_doc(text, filename="verslag.docx"):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
    print(f"Document opgeslagen als {filename}")

def transcribe_audio(file_path):
    segments, info = whisper_model.transcribe(file_path)

    text = ""
    for seg in segments:
        text += seg.text + " "

    return text.strip()

# =============================
# STATIC (CACHEABLE CONTENT)
# =============================
TEMPLATE = read_docx("template.docx")
BASISDOCUMENT = read_docx("basis.docx")
VOORBEELD1 = read_docx("voorbeeld1.docx")
VOORBEELD2 = read_docx("voorbeeld2.docx")
VOORBEELD3 = read_docx("voorbeeld3.docx")

SYSTEM_CONTEXT = f"""
Je bent een medische verslag-assistent.

TAKEN:
Je zet een consult transcript om in een gestructureerd verslag volgens een vaste template.

BELANGRIJK:
- Gebruik ALLEEN informatie uit: TRANSCRIPT, BASISDOCUMENT, TEMPLATE
- Voeg geen nieuwe informatie toe
- Verzin niets

TEMPLATE REGELS:
- Volg de structuur van het TEMPLATE exact
- Vul elk onderdeel in als “velden”
- Verander de opmaak of stijl niet

VOORBEELDEN:
- Gebruik voorbeelden alleen om te zien hoe het formaat eruit ziet
- NIET kopiëren van tekst uit voorbeelden

BASISDOCUMENT REGELS:
- Als een onderwerp (bv. supplementen, voeding, medicatie) voorkomt in het transcript:
  -> gebruik EXACT de info uit het BASISDOCUMENT
  -> herschrijf niet vrij
  -> kopieer inhoud zo letterlijk mogelijk

OUTPUT REGELS:
- Feitelijk en kort
- Geen extra uitleg
- Geen nieuwe secties maken
- Alles moet in het template passen
- Als informatie ontbreekt in transcript of basisdocument, schrijf ‘niet vermeld’ en verzin niets.
- Return alleen het ingevulde verslag, geen uitleg, geen opmerkingen.
- Gebruik bij het aanspreken altijd u of ander deftig taalgebruik nooit je

========================
TEMPLATE:
{TEMPLATE}

========================
BASIS ADVIEZEN:
{BASISDOCUMENT}

========================
VOORBEELD 1:
{VOORBEELD1}

========================
VOORBEELD 2:
{VOORBEELD2}

========================
VOORBEELD 3:
{VOORBEELD3}
"""

# =============================
# CORE FUNCTION
# =============================
def generate_report(transcript, notes=""):

    USER_CONTEXT = f"""
TRANSCRIPT:
{transcript}

NOTITIES:
{notes}
"""

    response = client.responses.create(
        model="gpt-5-mini",
        input=[
            {"role": "system", "content": SYSTEM_CONTEXT},
            {"role": "user", "content": USER_CONTEXT}
        ],
        temperature=0.1
    )

    return response.output[0].content[0].text

# =============================
# MAIN PROGRAM
# =============================
if __name__ == "__main__":

    print("=== Consult Verslag Generator ===")

    keuze = input("Gebruik (1) transcript of (2) audio mp3? ")

    if keuze == "1":
        transcript = input("Plak transcript:\n")
    elif keuze == "2":
        pad = input("Pad naar mp3 bestand: ")
        transcript = transcribe_audio(pad)
        print("\nTranscript geladen.\n")
    else:
        print("Ongeldige keuze")
        exit()

    notes = input("Notities (optioneel):\n")

    print("\nBezig met genereren...\n")
    verslag = generate_report(transcript, notes)

    print("\n=== RESULTAAT ===\n")
    print(verslag)

    if input("\nOpslaan als Word? (y/n): ").lower() == "y":
        save_to_doc(verslag)